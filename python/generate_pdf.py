"""
Generate PDF from DOCX with letterhead and signature.
Uses python-docx for template processing and docx2pdf for conversion.
"""

import os
import sys
import logging
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2pdf

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class PDFGenerator:
    """Class to generate PDFs from DOCX files with letterhead and signature."""
    
    def __init__(self, letterhead_template: str = None, signature_image: str = None):
        """
        Initialize PDF generator.
        
        Args:
            letterhead_template (str): Path to letterhead DOCX template
            signature_image (str): Path to signature image file
        """
        self.letterhead_template = letterhead_template
        self.signature_image = signature_image
    
    def generate_pdf(self, docx_path: str, output_path: str = None) -> str:
        """
        Generate PDF from DOCX file with letterhead and signature.
        
        Args:
            docx_path (str): Path to the input DOCX file
            output_path (str, optional): Path for the output PDF file
            
        Returns:
            str: Path to the generated PDF file
        """
        try:
            # Validate input file
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
            if not docx_path.lower().endswith('.docx'):
                raise ValueError("Input file must be a DOCX file")
            
            # Set output path if not provided
            if not output_path:
                output_path = os.path.splitext(docx_path)[0] + ".pdf"
            
            logger.info(f"Generating PDF from {docx_path}")
            
            # Load the document
            doc = Document(docx_path)
            
            # Add letterhead if template provided
            if self.letterhead_template and os.path.exists(self.letterhead_template):
                self._add_letterhead(doc)
            
            # Add signature if image provided
            if self.signature_image and os.path.exists(self.signature_image):
                self._add_signature(doc)
            
            # Add footer with timestamp
            self._add_footer(doc)
            
            # Save the modified document
            temp_docx = os.path.splitext(output_path)[0] + "_temp.docx"
            doc.save(temp_docx)
            
            # Convert DOCX to PDF
            docx2pdf.convert(temp_docx, output_path)
            
            # Clean up temporary file
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            
            logger.info(f"Successfully generated PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise
    
    def _add_letterhead(self, doc: Document):
        """Add letterhead to the document."""
        try:
            if not self.letterhead_template or not os.path.exists(self.letterhead_template):
                return
            
            # Load letterhead template
            letterhead_doc = Document(self.letterhead_template)
            
            # Copy header content from template
            if letterhead_doc.paragraphs:
                # Insert letterhead at the beginning
                for i, para in enumerate(letterhead_doc.paragraphs[:3]):  # First 3 paragraphs
                    if para.text.strip():
                        new_para = doc.paragraphs[0].insert_paragraph_before()
                        new_para.text = para.text
                        new_para.alignment = para.alignment
                        
                        # Copy formatting
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.font.size = run.font.size
                
                # Add some spacing
                doc.paragraphs[0].insert_paragraph_before()
                doc.paragraphs[0].insert_paragraph_before()
            
            logger.info("Added letterhead to document")
            
        except Exception as e:
            logger.error(f"Error adding letterhead: {str(e)}")
    
    def _add_signature(self, doc: Document):
        """Add signature image to the document."""
        try:
            if not self.signature_image or not os.path.exists(self.signature_image):
                return
            
            # Find the last paragraph or add a new one
            if not doc.paragraphs:
                doc.add_paragraph()
            
            # Add some spacing before signature
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Add signature paragraph
            sig_para = doc.add_paragraph()
            sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add signature image
            sig_para.add_run().add_picture(
                self.signature_image, 
                width=Inches(2.0)
            )
            
            # Add signature text
            sig_para.add_run("\n\nDr. Tony Forshaw\nEcho Guru Pty Ltd")
            
            logger.info("Added signature to document")
            
        except Exception as e:
            logger.error(f"Error adding signature: {str(e)}")
    
    def _add_footer(self, doc: Document):
        """Add footer with timestamp to the document."""
        try:
            # Add footer content
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            footer_para.add_run(f"Generated on {timestamp}")
            
            logger.info("Added footer to document")
            
        except Exception as e:
            logger.error(f"Error adding footer: {str(e)}")

def main():
    """Test function for PDF generator."""
    if len(sys.argv) < 2:
        print("Usage: python generate_pdf.py <docx_file_path> [output_path]")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # Initialize generator with optional assets
    assets_dir = os.path.join(os.path.dirname(__file__), '..', 'assets')
    letterhead = os.path.join(assets_dir, 'letterhead.docx')
    signature = os.path.join(assets_dir, 'signature.png')
    
    generator = PDFGenerator(
        letterhead_template=letterhead if os.path.exists(letterhead) else None,
        signature_image=signature if os.path.exists(signature) else None
    )
    
    try:
        result = generator.generate_pdf(docx_file, output_file)
        print(f"PDF generation successful: {result}")
    except Exception as e:
        print(f"PDF generation failed: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
